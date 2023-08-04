from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image, ExifTags, ImageFont, ImageDraw
import io
import os
import warnings

def rotate_image(image):
    try:
        exif = image._getexif()
        if exif is not None:
            for tag, value in exif.items():
                if tag in ExifTags.TAGS and ExifTags.TAGS[tag] == 'Orientation':
                    if value == 3:
                        return image.rotate(180, expand=True)
                    elif value == 6:
                        return image.rotate(270, expand=True)
                    elif value == 8:
                        return image.rotate(90, expand=True)
    except Exception as e:
        pass
    return image

def resize_image(image_path, scale_percent):
    image = Image.open(image_path)
    image = rotate_image(image)
    width, height = image.size

    # Verificar si la carpeta está dentro de las 61 primeras carpetas especiales
    folder_name = os.path.basename(os.path.dirname(image_path))
    special_folders = ["PE-16_" + str(i) for i in range(1, 62)]
    if folder_name in special_folders:
        # Escala del 24% para las 61 primeras carpetas
        scale_percent = 24

    new_width = int(width * scale_percent / 100)
    new_height = int(height * scale_percent / 100)
    resized_image = image.resize((new_width, new_height), Image.LANCZOS)
    return resized_image

def add_text_to_image(image, text, font_size):
    draw = ImageDraw.Draw(image)

    # Obtener la ruta de la fuente Arial Bold (arialbd.ttf) en el mismo directorio que el script
    font_path = os.path.join(os.path.dirname(__file__), "arialbd.ttf")
    font = ImageFont.truetype(font_path, font_size)

    text_width, text_height = draw.textsize(text, font=font)

    # Ajustar la posición del texto en la esquina inferior izquierda con un espacio del 1% del borde
    image_width, image_height = image.size
    margin = int(image_width * 0.01)
    x = margin
    y = image_height - text_height - margin

    # Agregar el texto en la imagen
    draw.text((x, y), text, fill="white", font=font)

    with warnings.catch_warnings():  # Usamos el contexto de manejo de advertencias
        warnings.simplefilter("ignore")  # Ignoramos todas las advertencias dentro del contexto
        text_width, text_height = draw.textsize(text, font=font)

    return image


def add_images_to_docx(doc, image_paths_group1, image_paths_group2, folder_name):
    labels = ["a) Vista terminal de inicio",
              "b) Vista terminal de fin",
              "c) Vista panorámica",
              "d) Vista relevante"]
    font_size = 12  # Tamaño de fuente actualizado

    # Agregar la primera tabla con las imágenes 1 y 2
    table1 = doc.add_table(rows=1, cols=2)
    table1.autofit = True

    for i, image_path in enumerate(image_paths_group1):
        image = resize_image(image_path, 8.3)  # Cambia la escala al 8.3%
        label_text = labels[i]
        image_with_text = add_text_to_image(image, label_text, font_size)

        cell = table1.cell(0, i)
        img_io = io.BytesIO()
        image_with_text.save(img_io, format='PNG')
        cell.paragraphs[0].add_run().add_picture(img_io, width=None, height=None)

    # Agregar la segunda tabla con las imágenes 3 y 4
    table2 = doc.add_table(rows=1, cols=2)
    table2.autofit = True

    for i, image_path in enumerate(image_paths_group2, start=2):
        image = resize_image(image_path, 8.3)  # Cambia la escala al 8.3%
        label_text = labels[i]
        image_with_text = add_text_to_image(image, label_text, font_size)

        cell = table2.cell(0, i - 2)
        img_io = io.BytesIO()
        image_with_text.save(img_io, format='PNG')
        cell.paragraphs[0].add_run().add_picture(img_io, width=None, height=None)

def get_image_paths(folder):
    image_paths = []
    for root, dirs, files in os.walk(folder):
        # Consideramos solo las carpetas que contengan exactamente 4 archivos de imagen (por ejemplo, 1.jpg, 2.jpg, 3.jpg, 4.jpg)
        if len(files) == 4 and all(file.lower().endswith((".jpg", ".jpeg")) for file in files):
            image_paths.extend([os.path.join(root, file) for file in files])
    return image_paths

def sort_folders(folder):
    prefix = folder.split('_')[0]
    if prefix == "PE-16":
        return (0, int(folder.split('_')[1]))
    elif prefix == "PE-3NE":
        return (1, int(folder.split('_')[1]))
    elif prefix == "AN-111":
        return (2, int(folder.split('_')[1]))
    else:
        return (3, 0)

def get_folders_with_four_photos(main_folder):
    folders_with_four_photos = []
    for folder in sorted(os.listdir(main_folder), key=sort_folders):
        folder_path = os.path.join(main_folder, folder)
        if os.path.isdir(folder_path):
            image_paths = get_image_paths(folder_path)
            if len(image_paths) == 4:
                folders_with_four_photos.append((folder_path, image_paths))
    return folders_with_four_photos

def main():
    main_folder = "D:\\FOTOS"  # Cambiar a la ruta de la carpeta principal
    folders_with_four_photos = get_folders_with_four_photos(main_folder)

    doc = Document()
    counter = 0  # Inicializar el contador

    with warnings.catch_warnings():  # Usamos el contexto de manejo de advertencias
        warnings.simplefilter("ignore")  # Ignoramos todas las advertencias dentro del contexto
        for folder_path, image_paths in folders_with_four_photos:
            counter += 1  # Incrementar el contador
            print("Cargando:", int(counter / len(folders_with_four_photos) * 100), "%")

            # Obtener el nombre de la carpeta de la tupla
            folder_name = os.path.basename(folder_path)

            image_paths_group1 = [os.path.join(folder_path, f"{i}.jpg") for i in range(1, 3)]
            image_paths_group2 = [os.path.join(folder_path, f"{i}.jpg") for i in range(3, 5)]

            if counter > 1:
                doc.add_page_break()  # Agregar salto de página después de procesar cada carpeta

            # Añadir texto encima de la primera tabla con el contador
            figure_text = f"Figura N°{counter}: Vistas del elemento de contención {folder_name.replace('_', ' ')}"
            paragraph = doc.add_paragraph(figure_text)
            paragraph_run = paragraph.runs[0]
            paragraph_run.bold = True
            paragraph_run.font.name = "Arial"
            paragraph_run.font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            add_images_to_docx(doc, image_paths_group1, image_paths_group2, folder_name)

    doc.save("D:\\FOTOS\\Panel-Fotografico-v4.docx")
    print("Documento de Word creado exitosamente.")

if __name__ == "__main__":
    main()
